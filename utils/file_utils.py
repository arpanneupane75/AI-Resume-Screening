# import fitz
# import docx
# from PIL import Image
# import pytesseract
# from docx import Document
# from fpdf import FPDF
# import io

# def extract_text_from_pdf(uploaded_file):
#     text = ""
#     try:
#         with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
#             for page in doc:
#                 page_text = page.get_text("text")
#                 if not page_text.strip():
#                     pix = page.get_pixmap()
#                     img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                     page_text = pytesseract.image_to_string(img)
#                 text += page_text + "\n"
#     except Exception as e:
#         return f"Error reading PDF: {e}"
#     return text.strip()

# def extract_text_from_docx(uploaded_file):
#     text = ""
#     try:
#         doc = docx.Document(uploaded_file)
#         text = "\n".join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         return f"Error reading DOCX: {e}"
#     return text.strip()

# def convert_docx_to_pdf(docx_text):
#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     for para in docx_text.split("\n"):
#         pdf.multi_cell(0, 10, para)
#     pdf_output = io.BytesIO()
#     pdf.output(pdf_output, "F")
#     return pdf_output.getvalue()
import fitz
import docx
from PIL import Image
import pytesseract
from fpdf import FPDF
import io
from typing import Union, List

# ---------------- PDF Extraction ----------------
def extract_text_from_pdf(file: Union[str, io.BytesIO]):
    """
    Extracts text from PDF. Uses OCR if PDF has no selectable text.
    """
    text = ""
    try:
        with fitz.open(stream=file.read() if hasattr(file, "read") else file, filetype="pdf") as doc:
            for page in doc:
                page_text = page.get_text("text")
                if not page_text.strip():  # Apply OCR if no text
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
    except Exception as e:
        raise RuntimeError(f"Error reading PDF: {e}")
    return text.strip()

# ---------------- DOCX Extraction ----------------
def extract_text_from_docx(file: Union[str, io.BytesIO]):
    """
    Extracts text from DOCX file.
    """
    text = ""
    try:
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX: {e}")
    return text.strip()

# ---------------- DOCX to PDF ----------------
def convert_docx_to_pdf(docx_text: str):
    """
    Converts a DOCX text string to a PDF in memory.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in docx_text.split("\n"):
        pdf.multi_cell(0, 10, para)
    pdf_output = io.BytesIO()
    pdf.output(pdf_output, "F")
    pdf_output.seek(0)
    return pdf_output.getvalue()

# ---------------- Batch Conversion ----------------
def batch_convert_docx_to_pdf(docx_texts: List[str]):
    """
    Converts a list of DOCX text strings to PDFs in memory.
    Returns a list of byte objects.
    """
    pdfs = []
    for text in docx_texts:
        pdfs.append(convert_docx_to_pdf(text))
    return pdfs
